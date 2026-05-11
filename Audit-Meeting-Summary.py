import streamlit as st
import os
import tempfile
import whisper
import ollama
import torch
import gc
import pandas as pd
from datetime import datetime
from pyannote.audio import Pipeline
from docx import Document
from fpdf import FPDF
import io
import json
import os
import sys
import shutil
import subprocess
import time

# 🚀 Cloud Ollama Startup (Runs only on first load)
if not os.environ.get("OLLAMA_SERVER_RUNNING"):
    st.sidebar.info("⏳ Starting Ollama server & downloading model... (1-2 min)")
    # Install Ollama if missing
    subprocess.run("curl -fsSL https://ollama.com/install.sh | sh", shell=True, check=True)
    # Start server in background
    subprocess.Popen(["ollama", "serve"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    os.environ["OLLAMA_SERVER_RUNNING"] = "1"
    # Wait for server to be ready
    time.sleep(5)
    # Pull model (cached in Cloud for ~24h)
    subprocess.run(["ollama", "pull", "llama3.2"], check=True)

# Set ffmpeg path before importing whisper/pyannote
if sys.platform == "win32":
    os.environ["FFMPEG_BINARY"] = r"C:\ffmpeg\bin\ffmpeg.exe"
    os.environ["FFPROBE_BINARY"] = r"C:\ffmpeg\bin\ffprobe.exe"

# ==================== PAGE CONFIG ====================
st.set_page_config(page_title="🎙️ Audio AI Pro", layout="wide")
st.title("🎙️ Audio AI Processor Pro")
st.markdown(
    "📁 Batch Upload →  Whisper Transcription → 🗣️ Pyannote Diarization → 🤖 Ollama Analysis → 📄 PDF/Word Export")

# ==================== SIDEBAR SETTINGS ====================
with st.sidebar:
    st.header("⚙️ Processing Settings")

    whisper_model = st.selectbox("Whisper Model", ["tiny", "base", "small", "medium", "large"], index=2,
                                 help="tiny=fastest, large=most accurate (requires ~3GB VRAM)")
    ollama_model = st.selectbox("Ollama Model", ["llama3.2", "mistral", "qwen2.5", "phi3"], index=0)
    detail_level = st.slider("Summary Detail Level", 1, 5, 3, help="1=Brief, 5=Highly Structured")

    st.divider()
    st.subheader("️ Speaker Diarization")
    enable_diarization = st.toggle("Enable Speaker Diarization", value=False,
                                   help="Requires Hugging Face token & ~2GB RAM")
    hf_token = st.text_input("Hugging Face Token (for Pyannote)", type="password",
                             value=st.session_state.get("hf_token", ""),
                             help="Get free token: https://huggingface.co/settings/tokens")
    if hf_token:
        st.session_state.hf_token = hf_token

    st.divider()
    st.subheader("📦 Export Settings")
    export_format = st.multiselect("Export Format", ["PDF", "Word"], default=["PDF"])

    st.divider()
    st.info(
        " 100% Local Processing\n• No data leaves your device\n• No API keys or charges\n• Requires `ffmpeg` installed")

    if st.button("🗑️ Clear GPU/RAM Cache", type="secondary"):
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        gc.collect()
        st.success("✅ Cache cleared!")


# ==================== MODEL LOADING ====================
@st.cache_resource
def load_whisper(model_name):
    return whisper.load_model(model_name, device="cuda" if torch.cuda.is_available() else "cpu")


@st.cache_resource
def load_diarization_pipeline(token):
    return Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=token)


# ==================== PROCESSING FUNCTIONS ====================
def transcribe_audio(model, file_path):
    # 🔧 FIX: Copy to a safe temp file with a clean name to bypass Windows ffmpeg path bugs
    ext = os.path.splitext(file_path)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        shutil.copy2(file_path, tmp.name)
        safe_path = tmp.name

    try:
        # Verify audio actually loads before transcribing
        audio = whisper.load_audio(safe_path)
        if len(audio) == 0:
            raise ValueError("Audio file is empty, corrupted, or in an unsupported format.")

        # Transcribe (force English for better accuracy & speed on podcasts)
        result = model.transcribe(safe_path, fp16=False, language="en")
        return result["text"].strip(), result["segments"]
    finally:
        # Clean up the safe copy
        if os.path.exists(safe_path):
            os.remove(safe_path)

def diarize_audio(pipeline, file_path, whisper_segments):
    diarization = pipeline(file_path)
    turns = [(turn.start, turn.end, turn[2]["speaker"]) for turn in diarization]

    mapped_segments = []
    for seg in whisper_segments:
        s, e = seg["start"], seg["end"]
        overlaps = [(spk, max(0, min(e, end) - max(s, start))) for start, end, spk in turns if
                    min(e, end) > max(s, start)]
        dominant = max(overlaps, key=lambda x: x[1])[0] if overlaps else "Unknown"
        mapped_segments.append({"speaker": dominant, "start": s, "end": e, "text": seg["text"].strip()})
    return mapped_segments


def ai_analysis(transcript, model_name, detail_level):
    detail_map = {1: "brief overview", 2: "concise summary", 3: "detailed breakdown", 4: "comprehensive analysis",
                  5: "highly structured & in-depth report"}

    prompts = {
        "summary": f"""Provide a {detail_map[detail_level]} of the transcript.
Structure:
 **Key Topics & Themes**
- ...
 **Critical Insights & Data Points**
- ...
✅ **Action Items & Takeaways**
- ...
Transcript: {transcript}""",

        "keywords": f"""Extract 10-15 key phrases, technical terms, and recurring topics from the transcript.
Output format:
- [Category]: Keyword1, Keyword2
- [Category]: Keyword3, Keyword4
Transcript: {transcript}""",

        "sentiment": f"""Analyze the overall sentiment, emotional tone shifts, and speaker attitudes in the transcript.
Structure:
 🌡️ **Overall Sentiment**: [Positive/Negative/Neutral/Mixed]
 📈 **Tone Progression**: [Beginning → Middle → End]
  **Key Emotional Triggers**: [List moments causing shifts]
 🎯 **Professional Tone Assessment**: [Formal/Casual/Assertive/Hesitant]
Transcript: {transcript}"""
    }

    results = {}
    for task, prompt in prompts.items():
        try:
            res = ollama.chat(model=model_name, messages=[{"role": "user", "content": prompt}],
                              options={"temperature": 0.3, "num_ctx": 8192})
            results[task] = res["message"]["content"]
        except Exception as e:
            results[task] = f"⚠️ Ollama Error: {str(e)}"
    return results


# ==================== EXPORT FUNCTIONS ====================
def generate_docx(data):
    doc = Document()
    doc.add_heading(f"Audio Analysis Report - {data['filename']}", 0)
    doc.add_paragraph(f"Processed: {data['timestamp']}")

    doc.add_heading("📝 Transcript", level=1)
    doc.add_paragraph(data["transcript"])

    if data.get("diarization"):
        doc.add_heading("🗣️ Speaker Diarization", level=1)
        for seg in data["diarization"]:
            doc.add_paragraph(f"[{seg['start']:.2f}s - {seg['end']:.2f}s] {seg['speaker']}: {seg['text']}",
                              style='List Bullet')

    for title, content in [("Summary", data["summary"]), ("Keywords", data["keywords"]),
                           ("Sentiment", data["sentiment"])]:
        doc.add_heading(f"🤖 {title}", level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, f"Audio Analysis Report - {data['filename']}")
    pdf.ln(5)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(0, 5, f"Processed: {data['timestamp']}")
    pdf.ln(5)

    for title, content in [("Transcript", data["transcript"]), ("Summary", data["summary"]),
                           ("Keywords", data["keywords"]), ("Sentiment", data["sentiment"])]:
        pdf.set_font("Helvetica", style="B", size=11)
        pdf.cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, content)
        pdf.ln(3)

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


# ==================== STREAMLIT UI ====================
uploaded_files = st.file_uploader("📁 Upload Audio Files", type=["mp3", "wav", "m4a", "flac", "ogg"],
                                  accept_multiple_files=True)

if uploaded_files:
    col1, col2 = st.columns([3, 1])
    col1.metric("Files Selected", len(uploaded_files))
    col2.metric("Total Size", f"{sum(f.size for f in uploaded_files) / 1024 / 1024:.2f} MB")

    if st.button(" Process All Files", type="primary", use_container_width=True):
        progress_bar = st.progress(0)
        status_text = st.empty()

        # Initialize session state for batch results
        if "batch_results" not in st.session_state:
            st.session_state.batch_results = {}

        whisper_model_obj = load_whisper(whisper_model)
        diar_pipeline = None
        if enable_diarization and hf_token:
            try:
                diar_pipeline = load_diarization_pipeline(hf_token)
            except Exception as e:
                st.error(f"❌ Failed to load diarization pipeline: {str(e)}")
                enable_diarization = False

        for idx, uploaded in enumerate(uploaded_files):
            status_text.text(f"Processing {idx + 1}/{len(uploaded_files)}: {uploaded.name}")

            file_ext = os.path.splitext(uploaded.name)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                tmp.write(uploaded.getvalue())
                tmp_path = tmp.name

            try:
                # 1. Transcribe
                transcript, segments = transcribe_audio(whisper_model_obj, tmp_path)

                # 2. Diarize (optional)
                diar_data = None
                if enable_diarization and diar_pipeline:
                    diar_data = diarize_audio(diar_pipeline, tmp_path, segments)

                # 3. AI Analysis
                ai_res = ai_analysis(transcript, ollama_model, detail_level)

                # 4. Store results
                st.session_state.batch_results[uploaded.name] = {
                    "filename": uploaded.name,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "transcript": transcript,
                    "segments": segments,
                    "diarization": diar_data,
                    "summary": ai_res["summary"],
                    "keywords": ai_res["keywords"],
                    "sentiment": ai_res["sentiment"]
                }

            except Exception as e:
                st.error(f" Failed to process {uploaded.name}: {str(e)}")
            finally:
                os.remove(tmp_path)
                gc.collect()

            progress_bar.progress((idx + 1) / len(uploaded_files))

        status_text.text("✅ All files processed!")
        progress_bar.empty()

# ==================== RESULTS DISPLAY ====================
if "batch_results" in st.session_state and st.session_state.batch_results:
    st.divider()
    st.subheader("📊 Processed Results")

    selected_file = st.selectbox("Select File to View", list(st.session_state.batch_results.keys()))
    data = st.session_state.batch_results[selected_file]

    tab_transcript, tab_diar, tab_summary, tab_keywords, tab_sentiment, tab_export = st.tabs([
        "📝 Transcript", "🗣️ Diarization", "📊 Summary", "🔑 Keywords", "💭 Sentiment", "📄 Export"
    ])

    with tab_transcript:
        st.text_area("Raw Transcript", data["transcript"], height=300, label_visibility="collapsed")

    with tab_diar:
        if data["diarization"]:
            df_diar = pd.DataFrame(data["diarization"])
            st.dataframe(df_diar, use_container_width=True)
            st.download_button(" Download Diarization CSV", df_diar.to_csv(index=False).encode(), "diarization.csv",
                               "text/csv")
        else:
            st.info("🗣️ Diarization was not enabled for this file.")

    with tab_summary:
        st.markdown(data["summary"])

    with tab_keywords:
        st.markdown(data["keywords"])

    with tab_sentiment:
        st.markdown(data["sentiment"])

    with tab_export:
        st.subheader("📦 Generate Export Files")
        col_a, col_b = st.columns(2)

        if "PDF" in export_format:
            pdf_buf = generate_pdf(data)
            col_a.download_button("📄 Download PDF", pdf_buf, f"{selected_file}_report.pdf", "application/pdf")

        if "Word" in export_format:
            docx_buf = generate_docx(data)
            col_b.download_button("📘 Download Word", docx_buf, f"{selected_file}_report.docx",
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.caption("Exports include transcript, diarization (if enabled), summary, keywords, and sentiment analysis.")
